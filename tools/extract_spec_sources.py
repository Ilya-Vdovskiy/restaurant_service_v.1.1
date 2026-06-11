from pathlib import Path

from docx import Document
from pypdf import PdfReader


SOURCES = {
    "template_pdf": Path(r"C:\Users\ilya0\Downloads\Шаблон спецификации (1).pdf"),
    "example_pdf": Path(r"C:\Users\ilya0\Downloads\Пример спецификации (1).pdf"),
    "spec_docx": Path(r"C:\Users\ilya0\Downloads\ЛР7. Спецификация (2).docx"),
    "test_plan_template": Path(r"C:\Users\ilya0\Downloads\Программа и методика испытаний (шаблон) (1).docx"),
    "explanatory_note": Path(
        r"C:\Users\ilya0\OneDrive\Desktop\3 курс\Цимбалюк\Курсач\пояснительная_записка_Вдовский_И_О_3093_исправ.docx"
    ),
}


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    chunks = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        chunks.append(f"\n\n===== PAGE {index} =====\n{fix_mojibake(text.strip())}")
    return "\n".join(chunks).strip()


def fix_mojibake(text: str) -> str:
    try:
        fixed = text.encode("cp1251").decode("utf-8")
    except UnicodeError:
        return text
    return fixed if fixed.count("�") <= text.count("�") else text


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    chunks = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            style = paragraph.style.name if paragraph.style else ""
            chunks.append(f"[{style}] {text}")

    for table_index, table in enumerate(doc.tables, start=1):
        chunks.append(f"\n[TABLE {table_index}]")
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            chunks.append(" | ".join(cells))

    return "\n".join(chunks).strip()


def main() -> None:
    out_dir = Path("analysis") / "spec_sources"
    out_dir.mkdir(parents=True, exist_ok=True)

    for name, path in SOURCES.items():
        if path.suffix.lower() == ".pdf":
            text = extract_pdf(path)
        else:
            text = extract_docx(path)

        output = out_dir / f"{name}.txt"
        output.write_text(text, encoding="utf-8")
        print(f"{name}: {len(text):,} chars -> {output}")


if __name__ == "__main__":
    main()
